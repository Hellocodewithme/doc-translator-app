import gradio as gr
from deep_translator import GoogleTranslator
from docx import Document
from pdf2docx import Converter
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import os

# Translate text
def translate_text(text, target_lang="hi"):
    return GoogleTranslator(source='auto', target=target_lang).translate(text)

# Translate DOCX
def translate_docx(input_docx_path, output_docx_path="translated.docx"):
    doc = Document(input_docx_path)
    for para in doc.paragraphs:
        if para.text.strip():
            para.text = translate_text(para.text)
    doc.save(output_docx_path)
    return output_docx_path

# Main function for Gradio
def translate_file(file):
    filename = file.name
    with open(filename, "wb") as f:
        f.write(file.read())

    output_file = "translated_hindi.docx"

    if filename.endswith(".docx"):
        output_file = translate_docx(filename, output_file)
    elif filename.endswith(".pdf"):
        try:
            cv = Converter(filename)
            docx_file = filename.replace(".pdf", ".docx")
            cv.convert(docx_file, start=0, end=None)
            cv.close()
            output_file = translate_docx(docx_file, output_file)
        except:
            images = convert_from_path(filename)
            doc = Document()
            for img in images:
                text = pytesseract.image_to_string(img)
                if text.strip():
                    doc.add_paragraph(translate_text(text))
            doc.save(output_file)
    elif filename.lower().endswith(("jpg", "jpeg", "png")):
        img = Image.open(filename)
        text = pytesseract.image_to_string(img)
        doc = Document()
        doc.add_paragraph(translate_text(text))
        doc.save(output_file)

    return output_file

# Gradio interface
iface = gr.Interface(
    fn=translate_file,
    inputs=gr.File(label="Upload PDF, DOCX, or Image"),
    outputs=gr.File(label="Download Translated DOCX"),
    title="📄 English → Hindi Translator"
)

iface.launch()
