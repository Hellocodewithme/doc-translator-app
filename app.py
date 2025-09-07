import streamlit as st
from deep_translator import GoogleTranslator
from docx import Document
from pdf2docx import Converter
from pdf2image import convert_from_path
from PIL import Image
import easyocr
import numpy as np

# Initialize EasyOCR reader
reader = easyocr.Reader(['en'])  # English OCR

# Page setup
st.set_page_config(page_title="Doc Translator", page_icon="🌐", layout="centered")
st.title("📄 English → Hindi Translator")
st.markdown("Upload a **PDF, Word, or Image file**, and download the Hindi translation.")

def translate_text(text, target_lang="hi"):
    return GoogleTranslator(source='auto', target=target_lang).translate(text)

def translate_docx(input_docx, output_docx="translated.docx", target_lang="hi"):
    doc = Document(input_docx)
    for para in doc.paragraphs:
        if para.text.strip():
            para.text = translate_text(para.text, target_lang)
    doc.save(output_docx)
    return output_docx

# File upload
uploaded_file = st.file_uploader("📂 Upload File", type=["pdf", "docx", "jpg", "png", "jpeg"])

if uploaded_file is not None:
    filename = uploaded_file.name
    with open(filename, "wb") as f:
        f.write(uploaded_file.getbuffer())

    output_file = "translated_hindi.docx"

    # DOCX
    if filename.endswith(".docx"):
        output_file = translate_docx(filename, output_file)

    # PDF
    elif filename.endswith(".pdf"):
        try:
            cv = Converter(filename)
            docx_file = filename.replace(".pdf", ".docx")
            cv.convert(docx_file, start=0, end=None)
            cv.close()
            output_file = translate_docx(docx_file, output_file)
        except:
            st.warning("⚠️ Using OCR for scanned PDF")
            images = convert_from_path(filename)
            doc = Document()
            for img in images:
                img_array = np.array(img)
                result = reader.readtext(img_array)
                text = " ".join([res[1] for res in result])
                if text.strip():
                    doc.add_paragraph(translate_text(text))
            doc.save(output_file)

    # Image
    elif filename.lower().endswith(("jpg", "jpeg", "png")):
        img = Image.open(filename)
        img_array = np.array(img)
        result = reader.readtext(img_array)
        text = " ".join([res[1] for res in result])
        doc = Document()
        doc.add_paragraph(translate_text(text))
        doc.save(output_file)

    # Download button
    with open(output_file, "rb") as f:
        st.download_button(
            "⬇️ Download Hindi Document",
            f,
            file_name="translated_hindi.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
